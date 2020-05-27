import docx
from docx.shared import Inches, Cm
import datetime
import os
import pandas as pd

# 폴더 생성
def create_folder(directory):
    try:
        if not os.path.exists(directory):
            os.makedirs(directory)
    except OSError:
        print(f'Error : {directory} 존재!')

student = input('학생명 :  ')
create_folder(student)


document = docx.Document()

# Title
document.add_heading(student+' 오답 시험지', 0)


while True:
    p = document.add_paragraph()
    r = p.add_run()

    question = input('틀린 문제는? :    ') # 세로 0열에 위치할 문제
    if question == 's': # # s키 입력 시, 종료
        break
    
    que_path = 'data/' + question + '.jpg'
    r.add_picture(que_path, width=Cm(6.8)) # 가로 길이 고정(세로는 비율에 따라)
    
    r.add_picture('boundary.png', width=Cm(1.0)) # 2단의 구분 경계선


    question2 = input('틀린 문제2는? :   ') # 세로 1열에 위치할 문제
    if question2 == 's': # s키 입력 시, 종료
        break

    que_path2 = 'data/' + question2 + '.jpg'
    r.add_picture(que_path2, width=Cm(6.8))

    
    p = document.add_paragraph('\n\n\n\n') # 여백(문제 풀이용 공간)


# 파일 저장
suffix = datetime.datetime.now().strftime('%y%m%d_%H%M%S')
filename = '_'.join([student, suffix])
print(f'파일 저장 : {student}/{filename}')

document.save(f'{student}/{filename}.docx')
